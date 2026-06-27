import os
from flask import Flask, request, render_template_string, redirect, url_for, session, Response, send_from_directory
from werkzeug.utils import secure_filename
import fitz  # PyMuPDF
import docx
import zipfile
import shutil
import google.generativeai as genai
import json
from PIL import Image
import pytesseract
import csv
from datetime import datetime
import io
import re
from difflib import SequenceMatcher

try:
    from sentence_transformers import SentenceTransformer
except ImportError:
    SentenceTransformer = None

try:
    import faiss
except ImportError:
    faiss = None

try:
    import numpy as np
except ImportError:
    np = None

# --- Gemini API Configuration ---
# The API key is read from an environment variable for security.
# Make sure to set `export GOOGLE_API_KEY="YOUR_API_KEY"` in your terminal.
# `GEMINI_API_KEY` is also accepted for backwards compatibility.
GEMINI_API_KEY = os.environ.get("GOOGLE_API_KEY") or os.environ.get("GEMINI_API_KEY")
genai.configure(api_key=GEMINI_API_KEY)
# Initialize the Flask application
app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
# Use a static secret key for development to ensure session persistence between restarts.
app.secret_key = "a-dev-secret-key-that-should-be-changed"
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# --- Helper Functions for Text Extraction ---

def gemini_api_key_error():
    if GEMINI_API_KEY:
        return None
    return {
        "error": "Gemini API key is missing. Set GOOGLE_API_KEY or GEMINI_API_KEY before running the app."
    }

_EMBEDDING_MODEL = None

def get_embedding_model():
    """Lazily loads the local embedding model used for evidence retrieval."""
    global _EMBEDDING_MODEL
    if _EMBEDDING_MODEL is None and SentenceTransformer:
        print("--- Loading sentence-transformers embedding model ---")
        try:
            _EMBEDDING_MODEL = SentenceTransformer("all-MiniLM-L6-v2")
        except Exception as e:
            print(f"--- Unable to load embedding model; evidence retrieval will use lexical fallback: {e} ---")
            return None
    return _EMBEDDING_MODEL

def extract_pdf_pages(file_path):
    """Extract page-level text from a PDF with OCR fallback for scanned pages."""
    pages = []
    try:
        with fitz.open(file_path) as doc:
            for page_num, page in enumerate(doc, start=1):
                pages.append({"page_number": page_num, "text": page.get_text()})

        total_text = "\n".join(page["text"] for page in pages).strip()
        if len(total_text) >= 100:
            return pages

        print("--- Minimal text found. Attempting page-level OCR fallback. ---")
        ocr_pages = []
        with fitz.open(file_path) as doc:
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                pix = page.get_pixmap(dpi=300)
                img = Image.open(io.BytesIO(pix.tobytes("png")))
                ocr_pages.append({
                    "page_number": page_num + 1,
                    "text": pytesseract.image_to_string(img)
                })
        return ocr_pages
    except Exception as e:
        return [{"page_number": 1, "text": f"Error reading PDF: {e}"}]

def extract_text_from_pdf(file_path):
    """Extracts text from a PDF file."""
    pages = extract_pdf_pages(file_path)
    return "\n".join(page.get("text", "") for page in pages)

def extract_text_from_docx(file_path):
    """Extracts text from a DOCX file."""
    try:
        doc = docx.Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text
    except Exception as e:
        return f"Error reading DOCX: {e}"

def extract_text_from_txt(file_path):
    """Extracts text from a TXT file."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        return f"Error reading TXT: {e}"

def build_page_chunks(uploaded_docs):
    """Build searchable page chunks for every uploaded PDF."""
    chunks = []
    for doc_key, doc_info in uploaded_docs.items():
        filename = doc_info.get('filename', '')
        if not filename.lower().endswith('.pdf'):
            continue

        doc_name = DOCUMENT_TYPES.get(doc_key, {}).get('name', doc_key)
        for page in extract_pdf_pages(doc_info['path']):
            text = page.get("text", "").strip()
            if not text:
                continue
            chunks.append({
                "doc_key": doc_key,
                "source_document": filename,
                "document_type": doc_name,
                "page_number": page.get("page_number"),
                "text": text
            })
    return chunks

def build_semantic_index(chunks):
    """Build a FAISS index for page chunks when embedding dependencies exist."""
    if not chunks or not SentenceTransformer or not faiss or np is None:
        return None

    model = get_embedding_model()
    if model is None:
        return None
    texts = [chunk["text"] for chunk in chunks]
    vectors = model.encode(texts, normalize_embeddings=True)
    vectors = np.asarray(vectors, dtype=np.float32)
    index = faiss.IndexFlatIP(vectors.shape[1])
    index.add(vectors)
    return {"index": index, "model": model}

def finding_to_query(finding):
    """Create the retrieval query for a finding without asking an LLM again."""
    query_parts = []
    for field in ("risk_title", "issue_summary", "short_summary", "recommended_action"):
        if finding.get(field):
            query_parts.append(str(finding[field]))
    for phrase in finding.get("evidence_queries", []) or []:
        if phrase:
            query_parts.append(str(phrase))
    return " ".join(query_parts).strip()

def normalize_for_match(text):
    return re.sub(r"\s+", " ", text or "").strip().lower()

def lexical_score(query, text):
    query_norm = normalize_for_match(query)
    text_norm = normalize_for_match(text)
    if not query_norm or not text_norm:
        return 0

    query_terms = set(re.findall(r"[a-z0-9]{3,}", query_norm))
    text_terms = set(re.findall(r"[a-z0-9]{3,}", text_norm))
    overlap = len(query_terms & text_terms) / max(len(query_terms), 1)
    ratio = SequenceMatcher(None, query_norm[:500], text_norm[:1000]).ratio()
    return overlap + ratio

def best_snippet(text, queries, max_chars=320):
    """Return the sentence/window most relevant to the query terms."""
    clean_text = re.sub(r"\s+", " ", text or "").strip()
    if len(clean_text) <= max_chars:
        return clean_text

    query_text = " ".join(q for q in queries if q)
    query_norm = normalize_for_match(query_text)
    for phrase in sorted([q for q in queries if q], key=len, reverse=True):
        phrase_norm = normalize_for_match(phrase)
        if len(phrase_norm) < 5:
            continue
        idx = normalize_for_match(clean_text).find(phrase_norm)
        if idx >= 0:
            start = max(0, idx - 120)
            end = min(len(clean_text), idx + len(phrase) + 180)
            return clean_text[start:end].strip()

    sentences = re.split(r"(?<=[.!?])\s+", clean_text)
    best = max(sentences, key=lambda sentence: lexical_score(query_norm, sentence), default=clean_text)
    if len(best) > max_chars:
        best = best[:max_chars].rsplit(" ", 1)[0]
    return best.strip()

def locate_evidence_for_finding(finding, chunks, semantic_index):
    """Find the best supporting PDF page for a finding using embeddings."""
    if not chunks:
        return None

    queries = [q for q in finding.get("evidence_queries", []) or [] if q]
    query = finding_to_query(finding)
    if not query:
        return None

    best_chunk = None
    if semantic_index:
        qvec = semantic_index["model"].encode([query], normalize_embeddings=True)
        qvec = np.asarray(qvec, dtype=np.float32)
        _, indices = semantic_index["index"].search(qvec, 1)
        if len(indices[0]) and indices[0][0] >= 0:
            best_chunk = chunks[int(indices[0][0])]
    else:
        print("--- Embedding dependencies unavailable; using lexical evidence fallback ---")
        best_chunk = max(chunks, key=lambda chunk: lexical_score(query, chunk["text"]), default=None)

    if not best_chunk:
        return None

    evidence_text = best_snippet(best_chunk["text"], queries + [query])
    return {
        "doc_key": best_chunk["doc_key"],
        "source_document": best_chunk["source_document"],
        "page_number": best_chunk["page_number"],
        "evidence_text": evidence_text
    }

def enrich_findings_with_evidence(key_findings, chunks):
    """Attach source_document, page_number, and evidence_text to each finding."""
    findings = key_findings.get("key_findings", []) if isinstance(key_findings, dict) else []
    if not findings:
        return key_findings

    semantic_index = build_semantic_index(chunks)
    for finding in findings:
        if "issue_summary" not in finding and finding.get("short_summary"):
            finding["issue_summary"] = finding["short_summary"]

        evidence = locate_evidence_for_finding(finding, chunks, semantic_index)
        if evidence:
            finding.update(evidence)
        else:
            finding.setdefault("source_document", None)
            finding.setdefault("page_number", None)
            finding.setdefault("evidence_text", "")
    return key_findings

# --- AI Interaction with Gemini ---
def analyze_document_with_ai(text):
    """
    Sends the document text to the Gemini API for analysis and returns the result.
    """
    missing_key = gemini_api_key_error()
    if missing_key:
        return missing_key
    if text.strip() == "Unsupported file type.":
        return {"error": "This file type is not supported for analysis."}
    if not text or not text.strip():
        return {"error": "The document appears to be empty or contains no readable text."}

    # This is the prompt that instructs the model on its task.
    prompt = f"""
You are an expert home inspector.

Your task is INFORMATION EXTRACTION.

IMPORTANT RULES:

1. Extract EVERY issue, defect, recommendation, hazard,
repair need, fungus condition, moisture issue, pest issue,
or condition likely to lead to future damage.


2. If the report mentions 10 separate issues, return 10 entries.

3. Treat each inspection item separately.

4. Include future-risk conditions even if no visible damage exists.

5. Include:
   - fungus
   - moisture
   - leaks
   - cracks
   - loose fixtures
   - pest activity
   - conditions likely to lead to future damage
   - recommendations for replacement or repair

6. Never combine multiple issues into one.

Return JSON only.

Example:

Input:
Item 8A: Garage door damaged by fungus.
Item 10A: Surface fungus on kitchen shelf.
Item 10B: Toilet loose.

Output:

{{
  "problems":[
    {{
      "location":"Garage",
      "item":"8A",
      "category":"Fungus",
      "description":"Garage side door and jambs damaged by fungus.",
      "severity":"High"
    }},
    {{
      "location":"Kitchen",
      "item":"10A",
      "category":"Fungus",
      "description":"Surface fungus found on kitchen shelf due to previous leaks.",
      "severity":"Medium"
    }},
    {{
      "location":"Hall Bathroom",
      "item":"10B",
      "category":"Plumbing",
      "description":"Toilet is loose or improperly mounted.",
      "severity":"Medium"
    }}
  ]
}}

Document:

{text}
"""
    try:
        print("--- Sending text to Gemini API for analysis ---")
        model = genai.GenerativeModel('gemini-3.1-flash-lite')
        generation_config = {
            "temperature": 0,
            "response_mime_type": "application/json"
        }

        response = model.generate_content(prompt, generation_config=generation_config)
        print("RAW RESPONSE (analyze_document_with_ai):", response.text)
        try:
            return json.loads(response.text)
        except json.JSONDecodeError:
            print(f"Error: AI did not return valid JSON. Response:\n{response.text}")
            return {"error": "AI response was not in a valid JSON format."}
    except Exception as e:
        print(f"Error calling Gemini API: {e}")
        return {"error": f"An error occurred: {e}"}

def summarize_document_with_ai(doc_name, text):
    """Generates a structured summary for a single document."""
    missing_key = gemini_api_key_error()
    if missing_key:
        return missing_key

    prompt = f"""
    You are a real estate analyst. Your task is to create a structured summary of the provided disclosure document.
    The document is a '{doc_name}'.

    Analyze the text below and extract the following information:
    - "purpose": A one-sentence description of the document's purpose.
    - "defects": An array of strings for explicitly mentioned problems (max 3 items).
    - "foundation": An array of strings for any mentions of foundation issues (max 3 items).
    - "systems": An array of strings for key property systems mentioned (max 3 items).
    - "other": An array of strings for other notable disclosures (max 3 items).

    Format your response as a single JSON object with keys: "purpose", "defects", "foundation", "systems", "other".
    Each key except "purpose" should contain an array of strings. If a section has no items, return an empty array.

    Document Text:
    ---
    {text}
    ---
    """
    try:
        print(f"--- Summarizing '{doc_name}' with Gemini ---")
        model = genai.GenerativeModel('gemini-3.1-flash-lite')
        generation_config = {"temperature": 0, "response_mime_type": "application/json"}
        response = model.generate_content(prompt, generation_config=generation_config)
        print("RAW RESPONSE (summarize_document_with_ai):", response.text)
        try:
            return json.loads(response.text)
        except json.JSONDecodeError:
            return {"error": "AI response for summary was not in a valid JSON format."}
    except Exception as e:
        print(f"Error during summarization: {e}")
        return {"error": f"An error occurred during summarization: {e}"}

def synthesize_findings_with_ai(all_docs_text):
    """Generates high-level insights from all documents combined."""
    missing_key = gemini_api_key_error()
    if missing_key:
        return missing_key

    prompt = f"""
    You are a senior real estate risk analyst. You have been given a collection of disclosure documents for a single property.
    Your primary task is to extract and synthesize as many material property risks as possible from the disclosure package.
    Do not limit the output to only the top risks. Include every meaningful risk, defect, hazard, repair concern, permit issue, water/moisture concern, pest issue, roof/foundation/system concern, legal/title concern, natural hazard, or recurring disclosure theme that could matter to a buyer.
    Prefer broad coverage over brevity, but do not duplicate the same risk. If the same issue appears in multiple documents, merge it into one finding and list the supporting documents.
    Sort findings by risk_level first (High, then Medium, then Low), then by likely buyer impact.
    Aim for 8-15 findings when the documents contain enough issues. Return fewer only if the documents genuinely contain fewer material risks.

    For each risk you identify, provide the following in a compact structured format:
    - "risk_title": A clear, concise title for the risk (e.g., "Water Intrusion History").
    - "risk_level": A string: "High", "Medium", or "Low".
    - "issue_summary": One short sentence, no more than 25 words.
    - "short_summary": The same value as "issue_summary" for backwards compatibility.
    - "evidence_sources": An array of strings listing the document types where this was mentioned (e.g., ["SPQ", "TDS"]).
    - "evidence_queries": An array of 1-3 short phrases copied directly from the source documents. Never paraphrase these phrases. They will be used later for embedding retrieval.
    - "potential_impacts": An array of strings describing potential consequences (max 2 short items, e.g., ["Mold growth", "Structural damage"]).
    - "cost_implication": A string: "High", "Medium", or "Low".
    - "recommended_action": One short actionable recommendation, no more than 18 words.
    - "next_steps": An array of strings with 1-2 short concrete next steps.

    Format your entire response as a single JSON object with two top-level keys: "key_findings" (an array of the risk objects) and "confidence" (a string: "High", "Medium", or "Low").
    If no significant cross-document risks are found, return an empty array.

    Combined Document Texts:
    ---
    {all_docs_text}
    ---
    """
    try:
        print("--- Synthesizing key findings with Gemini ---")
        # Use the more powerful model for this complex task
        model = genai.GenerativeModel('gemini-3.1-flash-lite')
        generation_config = {"temperature": 0, "response_mime_type": "application/json"}
        response = model.generate_content(prompt, generation_config=generation_config)
        print("RAW RESPONSE (synthesize_findings_with_ai):", response.text)
        try:
            return json.loads(response.text)
        except json.JSONDecodeError:
            return {"error": "AI response for synthesis was not in a valid JSON format."}
    except Exception as e:
        print(f"Error during synthesis: {e}")
        return {"error": f"An error occurred during synthesis: {e}"}


# --- Flask Routes ---

DOCUMENT_TYPES = {
    'tds': {'name': 'Transfer Disclosure Statement', 'required': True, 'description': 'Standard state-mandated disclosure form.'},
    'spq': {'name': 'Seller Property Questionnaire', 'required': True, 'description': 'Detailed questionnaire filled out by the seller.'},
    'nhd': {'name': 'Natural Hazard Disclosure Report', 'required': True, 'description': 'Report on natural hazards affecting the property.'},
    'prelim': {'name': 'Preliminary Report', 'required': True, 'description': 'Provides details on title, liens, and encumbrances.'},
    'inspection': {'name': 'Home Inspection Report', 'required': False, 'description': 'Professional home inspection findings.'},
    'non_foreign': {'name': "Seller's Affidavit of Non-Foreign Status", 'required': False, 'description': 'FIRPTA compliance form.'}
}

@app.route('/', methods=['GET', 'POST'])
def upload_file():
    if request.method == 'POST':
        if 'uploaded_docs' not in session:
            session['uploaded_docs'] = {}

        for doc_key, file in request.files.items():
            if file and file.filename:
                filename = secure_filename(file.filename)
                # Ensure a unique folder for this session/upload batch
                session_folder = session.get('session_folder')
                if not session_folder:
                    session_folder = datetime.now().strftime("%Y%m%d%H%M%S%f")
                    session['session_folder'] = session_folder
                
                save_dir = os.path.join(app.config['UPLOAD_FOLDER'], session_folder)
                os.makedirs(save_dir, exist_ok=True)
                file_path = os.path.join(save_dir, filename)
                file.save(file_path)

                session['uploaded_docs'][doc_key] = {
                    'filename': filename,
                    'uploaded_at': datetime.now().strftime("%Y-%m-%d %H:%M"),
                    'path': file_path,
                    'source': 'Manual'
                }
        
        # This is where you would trigger analysis. For now, we just show the checklist.
        return redirect(url_for('analyze_all'))

    # Initial page load
    # For a new session, clear old data
    if request.method == 'GET' and not request.args:
        session.pop('uploaded_docs', None)
        session.pop('session_folder', None)

    uploaded_docs = session.get('uploaded_docs', {})
    
    # Calculate progress statistics
    total_docs = len(DOCUMENT_TYPES)
    required_docs = {k:v for k,v in DOCUMENT_TYPES.items() if v['required']}
    
    loaded_count = len(uploaded_docs)
    required_loaded_count = sum(1 for k in required_docs if k in uploaded_docs)
    
    progress_percent = int((loaded_count / total_docs) * 100) if total_docs > 0 else 0
    all_required_loaded = required_loaded_count == len(required_docs)

    stats = {'total_docs': total_docs, 'loaded_count': loaded_count, 'progress_percent': progress_percent, 'all_required_loaded': all_required_loaded}

    return render_template_string(HTML_TEMPLATE, doc_types=DOCUMENT_TYPES, uploaded_docs=uploaded_docs, stats=stats)

@app.route('/analyze')
def analyze_all():
    # This is a placeholder for the analysis logic.
    # In a real app, this would iterate through session['uploaded_docs'],
    # run analysis on each, and display a results page.
    # For now, let's just show the uploaded files.
    uploaded_docs = session.get('uploaded_docs', {})
    if not uploaded_docs:
        return redirect(url_for('upload_file'))

    document_summaries = []
    all_text_for_synthesis = ""
    page_chunks = build_page_chunks(uploaded_docs)

    for doc_key, doc_info in uploaded_docs.items():
        text = process_single_file(doc_info['path'], doc_info['filename'])
        doc_name = DOCUMENT_TYPES[doc_key]['name']
        
        summary = summarize_document_with_ai(doc_name, text)
        document_summaries.append({
            'doc_key': doc_key,
            'doc_name': doc_name,
            'is_pdf': doc_info['filename'].lower().endswith('.pdf'),
            'summary': summary
        })
        
        all_text_for_synthesis += f"\n\n--- Start of Document: {doc_name} ---\n{text}\n--- End of Document: {doc_name} ---\n"

    # Generate the high-level synthesis
    key_findings = synthesize_findings_with_ai(all_text_for_synthesis)
    key_findings = enrich_findings_with_evidence(key_findings, page_chunks)

    return render_template_string(ANALYSIS_TEMPLATE, uploaded_docs=uploaded_docs, doc_types=DOCUMENT_TYPES, summaries=document_summaries, key_findings=key_findings)

@app.route('/bulk_analyze', methods=['POST'])
def bulk_analyze():
    """Handles the simple folder upload and immediate analysis."""
    uploaded_files = request.files.getlist("bulk_files")
    if not uploaded_files or uploaded_files[0].filename == '':
        return redirect(url_for('upload_file'))

    all_results = []
    # Create a unique folder for this bulk upload to avoid filename collisions
    session_folder = datetime.now().strftime("%Y%m%d%H%M%S%f")
    save_dir = os.path.join(app.config['UPLOAD_FOLDER'], session_folder)
    os.makedirs(save_dir, exist_ok=True)

    for file in uploaded_files:
        if file and file.filename:
            filename = secure_filename(file.filename)
            file_path = os.path.join(save_dir, filename)
            file.save(file_path)

            text = process_single_file(file_path, filename)
            analysis_result = analyze_document_with_ai(text)
            all_results.append({'filename': filename, 'result': analysis_result})

    # Save results to session for downloading
    session['analysis_results'] = all_results
    
    # Render the main template, showing the results but with an empty checklist
    stats = _get_progress_stats() # This will show 0% progress for the checklist
    return render_template_string(HTML_TEMPLATE, doc_types=DOCUMENT_TYPES, uploaded_docs={}, results=all_results, stats=stats)

@app.route('/view_doc/<doc_key>')
def view_doc(doc_key):
    """Serves an uploaded file for viewing."""
    uploaded_docs = session.get('uploaded_docs', {})
    doc_info = uploaded_docs.get(doc_key)
    if doc_info and os.path.exists(doc_info['path']):
        return send_from_directory(os.path.dirname(doc_info['path']), os.path.basename(doc_info['path']))
    return "File not found or session expired.", 404

@app.route('/pdf_viewer/<doc_key>')
def pdf_viewer(doc_key):
    """Renders a lightweight PDF.js viewer that can jump to and highlight evidence."""
    uploaded_docs = session.get('uploaded_docs', {})
    doc_info = uploaded_docs.get(doc_key)
    if not doc_info or not os.path.exists(doc_info['path']):
        return "File not found or session expired.", 404
    if not doc_info.get('filename', '').lower().endswith('.pdf'):
        return "PDF viewer is only available for PDF files.", 400

    return render_template_string(
        PDFJS_VIEWER_TEMPLATE,
        file_url=url_for('view_doc', doc_key=doc_key),
        initial_page=request.args.get('page', 1, type=int),
        highlight=request.args.get('highlight', ''),
        filename=doc_info.get('filename', 'Document')
    )

@app.route('/download_csv')
def download_csv():
    """Generates and serves a CSV file of the analysis results."""
    results = session.get('analysis_results', [])
    if not results:
        return redirect(url_for('upload_file'))

    # Use an in-memory string buffer
    output = io.StringIO()
    writer = csv.writer(output)

    # Write the header
    header = ['Source File', 'Location', 'Item', 'Category', 'Description', 'Severity']
    writer.writerow(header)

    # Write the data rows
    for item in results:
        filename = item.get('filename', 'N/A')
        if item.get('result', {}).get('problems'):
            for problem in item['result']['problems']:
                row = [
                    filename,
                    problem.get('location', 'N/A'),
                    problem.get('item', 'N/A'),
                    problem.get('category', 'N/A'),
                    problem.get('description', 'N/A'),
                    problem.get('severity', 'N/A')
                ]
                writer.writerow(row)

    output.seek(0)
    return Response(output, mimetype="text/csv", headers={"Content-Disposition":"attachment;filename=disclosure_analysis.csv"})

def _get_progress_stats():
    """Helper function to calculate and return upload progress stats."""
    uploaded_docs = session.get('uploaded_docs', {})
    total_docs = len(DOCUMENT_TYPES)
    required_docs = {k:v for k,v in DOCUMENT_TYPES.items() if v['required']}
    loaded_count = len(uploaded_docs)
    required_loaded_count = sum(1 for k in required_docs if k in uploaded_docs)
    progress_percent = int((loaded_count / total_docs) * 100) if total_docs > 0 else 0
    all_required_loaded = required_loaded_count == len(required_docs)
    return {'total_docs': total_docs, 'loaded_count': loaded_count, 'progress_percent': progress_percent, 'all_required_loaded': all_required_loaded}

def process_single_file(file_path, filename):
    """Helper function to extract text from a single file based on its extension."""
    text = ""
    if filename.lower().endswith('.pdf'):
        text = extract_text_from_pdf(file_path)
    elif filename.lower().endswith('.docx'):
        text = extract_text_from_docx(file_path)
    elif filename.lower().endswith('.txt'):
        text = extract_text_from_txt(file_path)
    else:
        # This message will be shown for unsupported files inside a zip
        # or for unsupported single uploads.
        text = "Unsupported file type."
    print(f"EXTRACTED TEXT LENGTH for {filename}: {len(text)}")
    if len(text.strip()) > 0:
        print(f"TEXT SNIPPET: {text.strip()[:500]}")
    return text

# --- HTML Template ---
# For simplicity, the HTML is included in the Python file.
# In a larger app, you would save this in a 'templates/index.html' file.
HTML_TEMPLATE = """
<!doctype html>
<html lang="en">
<head>
    <meta charset="utf-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Disclosure AI Uploader</title>
    <style>
        body { font-family: sans-serif; background-color: #f4f4f9; color: #333; margin: 40px; }
        .container { max-width: 800px; margin: auto; background: white; padding: 20px; border-radius: 8px; box-shadow: 0 0 10px rgba(0,0,0,0.1); }
        .progress-summary { text-align: center; margin-bottom: 30px; }
        .progress-bar { background-color: #e9ecef; border-radius: .25rem; height: 1rem; }
        .progress-bar-inner { background-color: #007bff; height: 1rem; border-radius: .25rem; }
        .section { border: 1px solid #dee2e6; border-radius: 5px; padding: 20px; margin-top: 30px; }
        .section h2 { margin-top: 0; }
        .section p { color: #6c757d; }
        .section input[type=text] { width: 100%; padding: 8px; margin-top: 10px; }
        .doc-table { width: 100%; margin-top: 20px; border-collapse: collapse; }
        .doc-table th, .doc-table td { padding: 12px; border-bottom: 1px solid #dee2e6; text-align: left; }
        .doc-table th { font-weight: bold; color: #495057; }
        .doc-list { list-style: none; padding: 0; }
        .doc-item { background: #f8f9fa; border: 1px solid #dee2e6; padding: 15px; margin-bottom: 10px; border-radius: 5px; display: flex; align-items: center; justify-content: space-between; }
        .doc-info h3 { margin: 0; font-size: 1.1em; }
        .doc-info p { margin: 5px 0 0; color: #6c757d; font-size: 0.9em; }
        .doc-status { text-align: right; }
        .status-uploaded { color: #28a745; font-weight: bold; }
        .status-missing { color: #dc3545; font-weight: bold; }
        .status-manual { color: #17a2b8; }
        .status-auto { color: #28a745; }
        .doc-actions .btn { padding: 5px 10px; font-size: 0.8em; cursor: pointer; }
        .btn-preview { background-color: #007bff; color: white; border: none; border-radius: 3px; }
        .btn-replace { background-color: #ffc107; color: black; border: none; border-radius: 3px; }
        h1, h2 { color: #0056b3; }
        .summary-box {
            display: flex;
            justify-content: space-around;
            background-color: #e9ecef;
            padding: 15px;
            border-radius: 5px;
            margin-bottom: 20px;
        }
        .summary-item { text-align: center; }
        .summary-item h3 { margin: 0; font-size: 24px; }
        .summary-item p { margin: 0; color: #6c757d; }
        .summary-high { color: #721c24 !important; }
        .summary-medium { color: #856404 !important; }
        table { width: 100%; border-collapse: collapse; margin-top: 20px; }
        th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }
        thead { background-color: #0056b3; color: white; }
        tr:nth-child(even) { background-color: #f2f2f2; }
        .severity-high { background-color: #f8d7da; color: #721c24; font-weight: bold; }
        .severity-medium { background-color: #fff3cd; color: #856404; font-weight: bold; }
        .severity-low { background-color: #d4edda; color: #155724; }
        .no-problems {
            padding: 15px;
            margin-top: 20px;
            background-color: #d4edda;
            color: #155724;
            border-left: 5px solid #28a745;
            font-weight: bold;
        }
        input[type=file], input[type=submit] { margin-top: 10px; padding: 8px; }
        .result { margin-top: 20px; padding: 15px; background: #e9ecef; border-left: 5px solid #0056b3; white-space: pre-wrap; }
        /* --- Loading Spinner Styles --- */
        #loader-overlay {
            position: fixed;
            top: 0;
            left: 0;
            width: 100%;
            height: 100%;
            background-color: rgba(0, 0, 0, 0.5);
            z-index: 9999;
            display: none; /* Hidden by default */
            justify-content: center;
            align-items: center;
            flex-direction: column;
            color: white;
        }
        .spinner {
            border: 8px solid #f3f3f3;
            border-top: 8px solid #0056b3;
            border-radius: 50%;
            width: 60px;
            height: 60px;
            animation: spin 1s linear infinite;
        }
        @keyframes spin { 0% { transform: rotate(0deg); } 100% { transform: rotate(360deg); } }
    </style>
</head>
<body>
    <div id="loader-overlay">
        <div class="spinner"></div>
        <p style="margin-top: 20px;">Analyzing documents, please wait...</p>
    </div>
    <div class="container">
        <h1>Upload Documents</h1>
        <p>Load property disclosure documents manually or use the disclosure link to auto-fetch available files.</p>

        <div class="progress-summary">
            <div class="progress-bar">
                <div class="progress-bar-inner" style="width: {{ stats.progress_percent }}%;"></div>
            </div>
            <p style="margin-top: 10px;"><b>{{ stats.progress_percent }}%</b></p>
            <p><b>{{ stats.loaded_count }} of {{ stats.total_docs }} Documents Loaded</b></p>
            <p>{{ stats.total_docs - stats.loaded_count }} documents remaining to analyze</p>
        </div>

        <div class="section">
            <h2>1. Auto Load from Disclosure Link</h2>
            <p>Paste the disclosure link below and we'll automatically fetch available documents.</p>
            <input type="text" placeholder="Paste disclosure link here...">
            <!-- In a real app, this would have JS to trigger a fetch -->
        </div>

        <div class="section">
            <h2>2. Manual Checklist Upload</h2>
            <p>Review, preview, or upload documents before running the analysis.</p>
            <form id="upload-form" method=post enctype=multipart/form-data>
                <table class="doc-table">
                    <thead>
                        <tr><th colspan="4">Required Documents</th></tr>
                        <tr>
                            <th>Document Type</th>
                            <th>Uploaded Document Name</th>
                            <th>Status</th>
                            <th>Actions</th>
                        </tr>
                    </thead>
                    <tbody>
                        {% for key, doc in doc_types.items() if doc.required %}
                        <tr>
                            <td><b>{{ doc.name }}*</b><br><small>{{ doc.description }}</small></td>
                            <td>{{ uploaded_docs[key].filename if uploaded_docs.get(key) else '—' }}</td>
                            <td>
                                {% if uploaded_docs.get(key) %}
                                    <span class="status-manual">Loaded (Manual)</span>
                                {% else %}
                                    <span class="status-missing">Missing</span>
                                {% endif %}
                            </td>
                            <td><input type="file" name="{{ key }}"></td>
                        </tr>
                        {% endfor %}
                    </tbody>
                    <thead>
                        <tr><th colspan="4" style="padding-top: 30px;">Optional Documents</th></tr>
                    </thead>
                    <tbody>
                        {% for key, doc in doc_types.items() if not doc.required %}
                        <tr>
                            <td><b>{{ doc.name }}</b><br><small>{{ doc.description }}</small></td>
                            <td>{{ uploaded_docs[key].filename if uploaded_docs.get(key) else '—' }}</td>
                            <td>
                                {% if uploaded_docs.get(key) %}
                                    <span class="status-manual">Loaded (Manual)</span>
                                {% else %}
                                    <span>Not uploaded</span>
                                {% endif %}
                            </td>
                            <td><input type="file" name="{{ key }}"></td>
                        </tr>
                        {% endfor %}
                    </tbody>
                </table>
                <input type=submit value="Upload Selected Files">
            </form>
        </div>

        <div class="section">
            <h2>3. Bulk Folder Upload & Analyze</h2>
            <p>Select a folder, and we will immediately analyze all documents inside it.</p>
            <form id="bulk-upload-form" action="/bulk_analyze" method="post" enctype="multipart/form-data">
                <input type="file" name="bulk_files" webkitdirectory multiple>
                <input type="submit" value="Upload Folder and Analyze">
            </form>
        </div>

    </div>
    <style>
        .download-button {
            display: inline-block;
            margin-top: 20px;
            padding: 10px 15px;
            background-color: #28a745;
            color: white;
            text-decoration: none;
            border-radius: 5px;
            font-weight: bold;
        }
    </style>
    <script>
        document.getElementById('upload-form').addEventListener('submit', function() {
            // Show the loader when the form is submitted
            const loader = document.getElementById('loader-overlay');
            loader.style.display = 'flex';
        });

        const bulkForm = document.getElementById('bulk-upload-form');
        if (bulkForm) {
            bulkForm.addEventListener('submit', function() {
                document.getElementById('loader-overlay').style.display = 'flex';
            });
        }

        const analyzeButton = document.getElementById('analyze-button');
        if (analyzeButton) {
            analyzeButton.addEventListener('click', function() {
                document.getElementById('loader-overlay').style.display = 'flex';
            });
        }
    </script>
</body>

</html>
"""

# --- Embedded PDF.js Viewer Template ---
PDFJS_VIEWER_TEMPLATE = """
<!doctype html>
<html lang="en">
<head>
    <meta charset="utf-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{{ filename }}</title>
    <link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/pdf.js/3.11.174/pdf_viewer.min.css">
    <style>
        html, body { margin: 0; height: 100%; background: #3f464f; font-family: sans-serif; }
        .toolbar {
            height: 44px;
            display: flex;
            align-items: center;
            gap: 10px;
            padding: 0 12px;
            background: #20242a;
            color: white;
            box-sizing: border-box;
        }
        .toolbar button {
            border: 1px solid #6b7280;
            background: #343a42;
            color: white;
            border-radius: 4px;
            padding: 5px 9px;
            cursor: pointer;
        }
        .toolbar span { font-size: 13px; }
        #viewer-shell {
            height: calc(100vh - 44px);
            overflow: auto;
            display: flex;
            justify-content: center;
            padding: 18px;
            box-sizing: border-box;
        }
        #page-wrap { position: relative; background: white; box-shadow: 0 2px 12px rgba(0,0,0,.35); }
        #pdf-canvas { display: block; }
        #text-layer {
            position: absolute;
            inset: 0;
            overflow: hidden;
            opacity: 1;
            line-height: 1;
        }
        #text-layer span {
            color: transparent;
            position: absolute;
            white-space: pre;
            transform-origin: 0% 0%;
        }
        #text-layer span.evidence-hit {
            background: rgba(255, 230, 0, .55);
            border-radius: 2px;
        }
        #status { margin-left: auto; color: #d1d5db; max-width: 45%; overflow: hidden; text-overflow: ellipsis; white-space: nowrap; }
    </style>
</head>
<body>
    <div class="toolbar">
        <button id="prev-page" title="Previous page">Prev</button>
        <button id="next-page" title="Next page">Next</button>
        <span>Page <strong id="page-num"></strong> / <span id="page-count">?</span></span>
        <span id="status">{{ filename }}</span>
    </div>
    <div id="viewer-shell">
        <div id="page-wrap">
            <canvas id="pdf-canvas"></canvas>
            <div id="text-layer" class="textLayer"></div>
        </div>
    </div>

    <script src="https://cdnjs.cloudflare.com/ajax/libs/pdf.js/3.11.174/pdf.min.js"></script>
    <script>
        const pdfUrl = {{ file_url | tojson }};
        const highlightText = {{ highlight | tojson }};
        let currentPage = Math.max(1, {{ initial_page | int }});
        let pdfDoc = null;
        let rendering = false;
        let pendingPage = null;
        const canvas = document.getElementById('pdf-canvas');
        const ctx = canvas.getContext('2d');
        const textLayer = document.getElementById('text-layer');
        const pageWrap = document.getElementById('page-wrap');

        pdfjsLib.GlobalWorkerOptions.workerSrc = 'https://cdnjs.cloudflare.com/ajax/libs/pdf.js/3.11.174/pdf.worker.min.js';

        function normalizeText(text) {
            return (text || '').toLowerCase().replace(/\\s+/g, ' ').trim();
        }

        function evidenceTerms(text) {
            const words = normalizeText(text).match(/[a-z0-9]{4,}/g) || [];
            const stop = new Set(['that', 'this', 'with', 'from', 'were', 'been', 'have', 'will', 'into', 'document', 'seller', 'property']);
            return [...new Set(words.filter(word => !stop.has(word)))].slice(0, 14);
        }

        function applyHighlight() {
            const fullNeedle = normalizeText(highlightText);
            const terms = evidenceTerms(highlightText);
            if (!fullNeedle && !terms.length) return;

            let firstHit = null;
            const spans = [...textLayer.querySelectorAll('span')];
            spans.forEach(span => {
                const spanText = normalizeText(span.textContent);
                const directHit = fullNeedle.length > 8 && fullNeedle.includes(spanText) && spanText.length > 3;
                const termHit = terms.some(term => spanText.includes(term));
                if (directHit || termHit) {
                    span.classList.add('evidence-hit');
                    firstHit = firstHit || span;
                }
            });

            if (firstHit) {
                firstHit.scrollIntoView({ block: 'center', inline: 'center' });
            }
        }

        async function renderPage(pageNumber) {
            rendering = true;
            const page = await pdfDoc.getPage(pageNumber);
            const shellWidth = document.getElementById('viewer-shell').clientWidth - 36;
            const baseViewport = page.getViewport({ scale: 1 });
            const scale = Math.min(1.55, Math.max(.75, shellWidth / baseViewport.width));
            const viewport = page.getViewport({ scale });

            canvas.width = viewport.width;
            canvas.height = viewport.height;
            pageWrap.style.width = `${viewport.width}px`;
            pageWrap.style.height = `${viewport.height}px`;
            textLayer.innerHTML = '';

            await page.render({ canvasContext: ctx, viewport }).promise;
            const textContent = await page.getTextContent();
            await pdfjsLib.renderTextLayer({
                textContentSource: textContent,
                container: textLayer,
                viewport,
                textDivs: []
            }).promise;

            document.getElementById('page-num').textContent = pageNumber;
            applyHighlight();
            rendering = false;

            if (pendingPage !== null) {
                const next = pendingPage;
                pendingPage = null;
                renderPage(next);
            }
        }

        function queueRender(pageNumber) {
            if (rendering) {
                pendingPage = pageNumber;
            } else {
                renderPage(pageNumber);
            }
        }

        document.getElementById('prev-page').addEventListener('click', () => {
            if (currentPage <= 1) return;
            currentPage -= 1;
            queueRender(currentPage);
        });

        document.getElementById('next-page').addEventListener('click', () => {
            if (!pdfDoc || currentPage >= pdfDoc.numPages) return;
            currentPage += 1;
            queueRender(currentPage);
        });

        pdfjsLib.getDocument(pdfUrl).promise.then(doc => {
            pdfDoc = doc;
            currentPage = Math.min(currentPage, pdfDoc.numPages);
            document.getElementById('page-count').textContent = pdfDoc.numPages;
            renderPage(currentPage);
        }).catch(error => {
            document.getElementById('status').textContent = `Unable to load PDF: ${error.message}`;
        });
    </script>
</body>
</html>
"""

# --- Analysis Page Template ---
ANALYSIS_TEMPLATE = """
<!doctype html>
<html lang="en">
<head>
    <meta charset="utf-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Disclosure AI - Analysis Report</title>
    <style>
        body { font-family: sans-serif; background-color: #f4f4f9; color: #333; margin: 32px; }
        .analysis-grid { display: grid; grid-template-columns: minmax(520px, 1.25fr) minmax(420px, .95fr); gap: 24px; align-items: start; max-width: 1600px; margin: auto; }
        .left-column { display: flex; flex-direction: column; gap: 25px; }
        .right-column { position: sticky; top: 20px; }
        .card { background: white; padding: 20px; border-radius: 8px; box-shadow: 0 4px 12px rgba(0,0,0,0.05); }
        h1, h2, h3 { color: #0056b3; }
        h2 { margin-top: 0; }
        #doc-selector { width: 100%; padding: 10px; margin-bottom: 15px; border-radius: 8px; border: 1px solid #ccc; }
        #pdf-viewer { width: 100%; height: 720px; border: none; border-radius: 8px; background: #20242a; }
        .summary-grid { display: grid; grid-template-columns: 1fr 1fr; gap: 20px; }
        .summary-card { background-color: #f8f9fa; padding: 15px; border-radius: 8px; }
        .summary-card h4 { margin-top: 0; border-bottom: 1px solid #e9ecef; padding-bottom: 8px; }
        .summary-card ul { padding-left: 20px; margin-bottom: 0; }
        .risk-card { background: #fff; border: 1px solid #f0caca; border-left: 5px solid #dc3545; border-radius: 8px; padding: 18px; margin-bottom: 16px; cursor: pointer; transition: border-color .15s ease, box-shadow .15s ease, transform .15s ease; }
        .risk-card:hover, .risk-card.active { border-color: #dc3545; box-shadow: 0 6px 18px rgba(220,53,69,.16); transform: translateY(-1px); }
        .risk-card h3 { margin-top: 0; display: flex; gap: 10px; align-items: center; justify-content: space-between; }
        .risk-card p { overflow-wrap: anywhere; }
        .risk-badge { padding: 3px 8px; border-radius: 12px; color: white; font-size: 0.8em; font-weight: bold; }
        .badge-high { background-color: #dc3545; }
        .badge-medium { background-color: #ffc107; }
        .badge-low { background-color: #28a745; }
        .evidence-box { background: #fffbea; border: 1px solid #f5df8e; border-radius: 6px; padding: 10px 12px; font-size: .9em; color: #4d3b00; }
        .evidence-meta { color: #6c757d; font-size: .8em; margin-bottom: 6px; }
        .error-box { background-color: #f8d7da; border: 1px solid #f5c6cb; color: #721c24; padding: 20px; border-radius: 5px; margin-top: 20px; }
        @media(max-width: 1000px) {
            .analysis-grid { grid-template-columns: 1fr; }
            .right-column { position: static; }
            #pdf-viewer { height: 560px; }
        }
    </style>
</head>
<body>
    <div class="analysis-grid">
        <div class="left-column">
            <div class="card">
                <h2>Document Viewer</h2>
                <select id="doc-selector">
                    {% for summary_item in summaries %}
                    <option value="{{ summary_item.doc_key }}" data-is-pdf="{{ 'true' if summary_item.is_pdf else 'false' }}">{{ summary_item.doc_name }}</option>
                    {% endfor %}
                </select>
                <iframe id="pdf-viewer" title="PDF evidence viewer"></iframe>
            </div>
            <div class="card" id="summary-section">
                <!-- This section will be populated by JavaScript -->
            </div>
        </div>

        <div class="right-column">
            <div class="card">
                <h2>Key Findings & AI Insights</h2>
                <p>Overall Analysis Confidence: <strong>{{ key_findings.get('confidence', 'N/A') }}</strong></p>
                {% if key_findings.get('error') %}
                    <div class="error-box"><p><strong>Error during synthesis:</strong> {{ key_findings.get('error') }}</p></div>
                {% endif %}
                {% for finding in key_findings.get('key_findings', []) %}
                    <div class="risk-card" data-finding-index="{{ loop.index0 }}" role="button" tabindex="0">
                        <h3>
                            {{ finding.risk_title }}
                            <span class="risk-badge badge-{{ finding.risk_level|lower }}">{{ finding.risk_level }}</span>
                        </h3>
                        <p><em>{{ finding.issue_summary or finding.short_summary }}</em></p>
                        {% if finding.source_document and finding.page_number %}
                            <div class="evidence-box">
                                <div class="evidence-meta">{{ finding.source_document }} · Page {{ finding.page_number }}</div>
                                {{ finding.evidence_text }}
                            </div>
                        {% else %}
                            <div class="evidence-box">
                                <div class="evidence-meta">Evidence location unavailable</div>
                                Install embedding dependencies and re-run analysis for semantic page retrieval.
                            </div>
                        {% endif %}
                        <p style="font-size: 0.8em; color: #6c757d; margin-top: 12px;">
                            <strong>Sources:</strong> {{ finding.evidence_sources|join(', ') }} |
                            <strong>Cost:</strong> {{ finding.cost_implication }} |
                            <strong>Action:</strong> {{ finding.recommended_action }}
                        </p>
                        {% if finding.potential_impacts %}
                            <p style="font-size: 0.85em; margin-bottom: 0;"><strong>Impacts:</strong> {{ finding.potential_impacts|join(', ') }}</p>
                        {% endif %}
                        {% if finding.next_steps %}
                            <p style="font-size: 0.85em; margin-top: 6px;"><strong>Next:</strong> {{ finding.next_steps|join(', ') }}</p>
                        {% endif %}
                    </div>
                {% else %}
                    <p>No high-priority cross-document risks were identified.</p>
                {% endfor %}
            </div>
            <div class="card">
                <h2>AI Follow-up</h2>
                <p style="color: #6c757d;">Ask the AI a follow-up question about these findings.</p>
                <input type="text" placeholder="e.g., 'What are the typical costs to repair foundation cracks?'" style="width: 95%; padding: 10px; border: 1px solid #ccc; border-radius: 5px;">
            </div>
        </div>
    </div>
    <script>
        const summariesData = {{ summaries | tojson }};
        const findingsData = {{ key_findings.get('key_findings', []) | tojson }};
        const docSelector = document.getElementById('doc-selector');
        const pdfViewer = document.getElementById('pdf-viewer');
        const summarySection = document.getElementById('summary-section');
        const riskCards = [...document.querySelectorAll('.risk-card')];

        function escapeHtml(value) {
            return String(value || '').replace(/[&<>"']/g, char => ({
                '&': '&amp;',
                '<': '&lt;',
                '>': '&gt;',
                '"': '&quot;',
                "'": '&#39;'
            }[char]));
        }

        function renderSummary(docKey) {
            const summaryItem = summariesData.find(s => s.doc_key === docKey);
            if (!summaryItem) {
                summarySection.innerHTML = '<h2>Summary Not Found</h2>';
                return;
            }

            const summary = summaryItem.summary;
            let html = `<h2>Summary: ${escapeHtml(summaryItem.doc_name)}</h2>`;
            html += `<p><em>${escapeHtml(summary.purpose || 'N/A')}</em></p>`;
            
            if (summary.error) {
                html += `<div class="error-box"><p><strong>Error:</strong> ${escapeHtml(summary.error)}</p></div>`;
            } else {
                html += '<div class="summary-grid">';
                if (summary.defects && summary.defects.length > 0) {
                    html += '<div class="summary-card"><h4>Known Defects</h4><ul>' + summary.defects.map(i => `<li>${escapeHtml(i)}</li>`).join('') + '</ul></div>';
                }
                if (summary.foundation && summary.foundation.length > 0) {
                    html += '<div class="summary-card"><h4>Foundation Risks</h4><ul>' + summary.foundation.map(i => `<li>${escapeHtml(i)}</li>`).join('') + '</ul></div>';
                }
                if (summary.systems && summary.systems.length > 0) {
                    html += '<div class="summary-card"><h4>Property Systems</h4><ul>' + summary.systems.map(i => `<li>${escapeHtml(i)}</li>`).join('') + '</ul></div>';
                }
                if (summary.other && summary.other.length > 0) {
                    html += '<div class="summary-card"><h4>Other Disclosures</h4><ul>' + summary.other.map(i => `<li>${escapeHtml(i)}</li>`).join('') + '</ul></div>';
                }
                html += '</div>';
            }
            summarySection.innerHTML = html;
        }

        function loadPdf(docKey, page = 1, evidenceText = '') {
            const params = new URLSearchParams({
                page: page || 1,
                highlight: evidenceText || ''
            });
            pdfViewer.src = `/pdf_viewer/${encodeURIComponent(docKey)}?${params.toString()}`;
            docSelector.value = docKey;
            renderSummary(docKey);
        }

        function chooseDefaultFinding() {
            const severityRank = { High: 0, Medium: 1, Low: 2 };
            return findingsData
                .map((finding, index) => ({ finding, index }))
                .filter(item => item.finding.doc_key && item.finding.page_number)
                .sort((a, b) => {
                    const aRank = severityRank[a.finding.risk_level] ?? 3;
                    const bRank = severityRank[b.finding.risk_level] ?? 3;
                    return aRank - bRank || a.index - b.index;
                })[0];
        }

        function selectFinding(index) {
            const finding = findingsData[index];
            if (!finding || !finding.doc_key || !finding.page_number) return;

            riskCards.forEach(card => card.classList.toggle('active', Number(card.dataset.findingIndex) === index));
            loadPdf(finding.doc_key, finding.page_number, finding.evidence_text || '');
        }

        docSelector.addEventListener('change', (event) => {
            const selectedOption = event.target.selectedOptions[0];
            const selectedDocKey = event.target.value;
            riskCards.forEach(card => card.classList.remove('active'));
            renderSummary(selectedDocKey);

            if (selectedOption && selectedOption.dataset.isPdf === 'true') {
                loadPdf(selectedDocKey);
            } else {
                pdfViewer.removeAttribute('src');
            }
        });

        riskCards.forEach(card => {
            const index = Number(card.dataset.findingIndex);
            card.addEventListener('click', () => selectFinding(index));
            card.addEventListener('keydown', event => {
                if (event.key === 'Enter' || event.key === ' ') {
                    event.preventDefault();
                    selectFinding(index);
                }
            });
        });

        // Initial render
        if (summariesData.length > 0) {
            const defaultFinding = chooseDefaultFinding();
            if (defaultFinding) {
                selectFinding(defaultFinding.index);
            } else {
                const firstPdf = summariesData.find(item => item.is_pdf) || summariesData[0];
                renderSummary(firstPdf.doc_key);
                if (firstPdf.is_pdf) {
                    loadPdf(firstPdf.doc_key);
                }
            }
        }
    </script>
</body>
</html>
"""

if __name__ == '__main__':
    # Runs the web server. Access it at http://127.0.0.1:5000
    app.run(host='0.0.0.0', port=5001, debug=True)
